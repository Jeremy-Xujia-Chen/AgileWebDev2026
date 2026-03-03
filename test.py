#!/usr/bin/env python3
"""
把 Word 文档 (.docx) 里所有包含 git 命令的文本改成绿色。
用法: python color_git_commands.py [文档路径]
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)

# 绿色 (常见终端绿)
GREEN = RGBColor(0x00, 0x80, 0x00)

def color_git_in_docx(docx_path: str) -> None:
    path = Path(docx_path)
    if not path.exists():
        print(f"文件不存在: {path}")
        sys.exit(1)
    if path.suffix.lower() != ".docx":
        print("只支持 .docx 文件")
        sys.exit(1)

    doc = Document(str(path))
    count = 0

    for para in doc.paragraphs:
        for run in para.runs:
            # 包含 "git " 的 run 视为 git 命令行，改成绿色
            if "git " in run.text or run.text.strip().startswith("git "):
                run.font.color.rgb = GREEN
                count += 1

    # 也检查表格里的单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "git " in run.text or run.text.strip().startswith("git "):
                            run.font.color.rgb = GREEN
                            count += 1

    out_path = path.parent / f"{path.stem}_green{path.suffix}"
    doc.save(str(out_path))
    print(f"已处理 {count} 处 git 命令，另存为: {out_path}")


if __name__ == "__main__":
    default = Path(__file__).resolve().parent.parent.parent / "Desktop" / "5505 ppt" / "01Agile-1..docx"
    docx_path = sys.argv[1] if len(sys.argv) > 1 else str(default)
    color_git_in_docx(docx_path)
